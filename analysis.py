import os
from pathlib import Path

from langchain_community.vectorstores import FAISS
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_text_splitters import CharacterTextSplitter
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.messages import HumanMessage
from langchain_nvidia_ai_endpoints import ChatNVIDIA, NVIDIAEmbeddings
import gradio as gr
import base64
from IPython.display import HTML, display

nvapi_key = "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"
os.environ["NVIDIA_API_KEY"] = nvapi_key

_VISION_MODEL_PRIMARY = os.environ.get(
    "NVIDIA_VISION_MODEL", "microsoft/phi-3.5-vision-instruct"
)
_VISION_MODEL_FALLBACK = "meta/llama-3.2-11b-vision-instruct"

_EMBEDDING_MODEL = os.environ.get(
    "NVIDIA_EMBEDDING_MODEL", "nvidia/nv-embedqa-e5-v5"
)

def _faiss_parent_dir() -> str:
    preferred = "zh_data"
    if os.path.exists(preferred) and not os.path.isdir(preferred):
        return "analysis_rag_index"
    return preferred


_FAISS_INDEX_DIR = os.path.join(
    _faiss_parent_dir(),
    "faiss_" + _EMBEDDING_MODEL.replace("/", "_").replace("-", "_"),
)

_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tiff", ".tif"}
_DOCX_EXTS = {".docx"}


class PdfUnreadableError(Exception):
    """ PDF cannot be read (e.g., encrypted, corrupted). """

class DocxUnreadableError(Exception):
    """DOCX cannot be read (e.g., corrupted)."""

def _extract_pdf_text_units(path: str) -> tuple[list[str], list[str]]:
    """ Extract the PDF text page by page, returning (a list of paragraphs, with source tags for each paragraph). """
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError(
            "To analyze PDFs, you need to install pypdf. Please execute: `pip install pypdf`"
        ) from exc

    reader = PdfReader(path)
    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception:
            pass
        if getattr(reader, "is_encrypted", False):
            raise PdfUnreadableError(
                "This PDF is encrypted and cannot be opened with a blank password. Please upload an unencrypted file or decrypt it locally first."
            )

    base = os.path.basename(path)
    data: list[str] = []
    sources: list[str] = []
    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if text:
            data.append(text)
            sources.append(f"{base} page {i + 1}")
    return data, sources


def _extract_docx_text_units(path: str) -> tuple[list[str], list[str]]:
    """Extract DOCX text, returning (paragraphs, source tags)."""
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "To analyze DOCX files, you need to install python-docx. "
            "Please execute: `pip install python-docx`"
        ) from exc

    try:
        doc = Document(path)
    except Exception as exc:
        raise DocxUnreadableError(
            "This DOCX file cannot be opened (it may be corrupted)."
        ) from exc

    base = os.path.basename(path)
    data: list[str] = []
    sources: list[str] = []

    for i, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip()
        if text:
            data.append(text)
            sources.append(f"{base} paragraph {i + 1}")

    # Also include table cell text if present (common in DOCX reports)
    try:
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    text = (cell.text or "").strip()
                    if text:
                        data.append(text)
                        sources.append(f"{base} table{ti + 1} r{ri + 1} c{ci + 1}")
    except Exception:
        # Tables are optional; don't fail extraction if some table is malformed
        pass

    return data, sources


def _run_rag_from_text_units(data: list[str], sources: list[str], user_prompt: str) -> tuple[str, str]:
    documents = [d for d in data if d != ""]
    if not documents:
        return "No usable text was parsed from the file, so a search cannot be established. ", "<ul><li>(no text)</li></ul>"

    llm = ChatNVIDIA(
        model="microsoft/phi-3-small-128k-instruct",
        nvidia_api_key=nvapi_key,
        max_completion_tokens=512,
    )
    first = llm.invoke(user_prompt)
    preview = getattr(first, "content", None)
    if preview is None:
        preview = str(first)
    html = f"<ul><li>{preview}</li></ul>"

    embedder = NVIDIAEmbeddings(
        model=_EMBEDDING_MODEL,
        nvidia_api_key=nvapi_key,
    )
    text_splitter = CharacterTextSplitter(chunk_size=400, separator=" ")
    docs: list[str] = []
    metadatas: list[dict] = []

    for i, d in enumerate(documents):
        splits = text_splitter.split_text(d)
        docs.extend(splits)
        metadatas.extend([{"source": sources[i]}] * len(splits))

    Path(_FAISS_INDEX_DIR).parent.mkdir(parents=True, exist_ok=True)
    store = FAISS.from_texts(docs, embedder, metadatas=metadatas)
    store.save_local(_FAISS_INDEX_DIR)
    store = FAISS.load_local(
        _FAISS_INDEX_DIR, embedder, allow_dangerous_deserialization=True
    )
    retriever = store.as_retriever()

    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "Answer solely based on the following context:\n<Documents>\n{context}\n</Documents>",
            ),
            ("user", "{question}"),
        ]
    )
    chain = (
        {"context": retriever, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )
    text = chain.invoke(user_prompt)
    return text, html


def _resolve_upload_path(file) -> str | None:
    """ Gradio may accept str, Path, NamedString, or a list of single files. """
    if file is None:
        return None
    if isinstance(file, (list, tuple)):
        file = file[0] if file else None
    if file is None:
        return None
    path = getattr(file, "name", file)
    return str(path)


def process_text(file, user_prompt):
    data: list[str] = []
    sources: list[str] = []
    if file.endswith(".txt"):
        with open(file, encoding="utf-8") as f:
            for line in f.readlines():
                if line:
                    data.append(line)
                    sources.append(file)
    return _run_rag_from_text_units(data, sources, user_prompt)


def process_pdf(file, user_prompt):
    data, sources = _extract_pdf_text_units(file)
    return _run_rag_from_text_units(data, sources, user_prompt)

def process_docx(file, user_prompt):
    data, sources = _extract_docx_text_units(file)
    return _run_rag_from_text_units(data, sources, user_prompt)


def image2b64(image_file):
    with open(image_file, "rb") as f:
        image_b64 = base64.b64encode(f.read()).decode()
        return image_b64


def _image_mime_type(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext in (".jpg", ".jpeg"):
        return "image/jpeg"
    if ext == ".webp":
        return "image/webp"
    if ext == ".gif":
        return "image/gif"
    if ext in (".bmp",):
        return "image/bmp"
    if ext in (".tiff", ".tif"):
        return "image/tiff"
    return "image/png"


def display_image(image_path):
    mime = _image_mime_type(image_path)
    with open(image_path, "rb") as image_file:
        encoded_string = base64.b64encode(image_file.read()).decode()
    html = f'<img src="data:{mime};base64,{encoded_string}" />'
    return html


def process_image(file, user_prompt):
    image_b64 = image2b64(file)
    mime = _image_mime_type(file)
    messages = [
        HumanMessage(
            content=[
                {"type": "text", "text": user_prompt or "Please describe this image."},
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime};base64,{image_b64}"},
                },
            ]
        )
    ]
    last_error: Exception | None = None
    for model_name in (_VISION_MODEL_PRIMARY, _VISION_MODEL_FALLBACK):
        try:
            chart_reading = ChatNVIDIA(
                model=model_name,
                nvidia_api_key=nvapi_key,
                max_completion_tokens=1024,
            )
            result = chart_reading.invoke(messages)
            return result.content
        except Exception as exc:
            last_error = exc
            continue
    assert last_error is not None
    raise last_error


def big_model_output(file, user_prompt):
    path = _resolve_upload_path(file)
    if not path:
        return "Please upload the file first.", ""
    file_extension = os.path.splitext(path)[1].lower()

    if file_extension == ".txt":
        return process_text(path, user_prompt)
    if file_extension == ".pdf":
        try:
            return process_pdf(path, user_prompt)
        except ImportError as e:
            return str(e), ""
        except PdfUnreadableError as e:
            return str(e), ""
    if file_extension in _DOCX_EXTS:
        try:
            return process_docx(path, user_prompt)
        except ImportError as e:
            return str(e), ""
        except DocxUnreadableError as e:
            return str(e), ""
    if file_extension == ".doc":
        return (
            "Legacy .doc is not supported directly. Please convert it to .docx or .pdf and upload again.",
            "",
        )
    if file_extension in _IMAGE_EXTS:
        image_html = display_image(path)
        return process_image(path, user_prompt), image_html
    return (
        f"Unsupported file types: {file_extension or ' (No file extension) '}.",
        "",
    )


iface = gr.Interface(
    fn=big_model_output,
    inputs=[
        gr.File(),
        gr.Textbox(lines=1)
    ],
    outputs=[
        "text",
        gr.HTML()
    ],
    title="RAG Analysis Tools with NVIDIA AI Agent",
    description = "Upload a txt, pdf or image file, enter your question, and we will analyze it."
)

iface.launch()